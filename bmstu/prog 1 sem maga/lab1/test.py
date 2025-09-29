import os
import re
import xml.etree.ElementTree as ET
from collections import Counter, defaultdict
from bs4 import BeautifulSoup
import openpyxl
from docx import Document

class DocumentVectorizer:
    def __init__(self):
        self.vocabulary = set()
        self.word_to_index = {}
        self.index_to_word = {}
        self.stop_words = {
            'и', 'в', 'во', 'не', 'что', 'он', 'на', 'я', 'с', 'со', 'как', 'а', 
            'то', 'все', 'она', 'так', 'его', 'но', 'да', 'ты', 'к', 'у', 'же', 
            'вы', 'за', 'бы', 'по', 'только', 'ее', 'мне', 'было', 'вот', 'от', 
            'меня', 'еще', 'нет', 'о', 'из', 'ему', 'теперь', 'когда', 'даже', 
            'ну', 'вдруг', 'ли', 'если', 'уже', 'или', 'ни', 'быть', 'был', 
            'него', 'до', 'вас', 'нибудь', 'опять', 'уж', 'вам', 'ведь', 'там', 
            'потом', 'себя', 'ничего', 'ей', 'может', 'они', 'тут', 'где', 'есть', 
            'надо', 'ней', 'для', 'мы', 'тебя', 'их', 'чем', 'была', 'сам', 'чтоб', 
            'без', 'будто', 'чего', 'раз', 'тоже', 'себе', 'под', 'будет', 'ж', 
            'тогда', 'кто', 'этот', 'того', 'потому', 'этого', 'какой', 'совсем', 
            'ним', 'здесь', 'этом', 'один', 'почти', 'мой', 'тем', 'чтобы', 'нее', 
            'сейчас', 'были', 'куда', 'зачем', 'всех', 'никогда', 'можно', 'при', 
            'наконец', 'два', 'об', 'другой', 'хоть', 'после', 'над', 'больше', 
            'тот', 'через', 'эти', 'нас', 'про', 'всего', 'них', 'какая', 'много', 
            'разве', 'три', 'эту', 'моя', 'впрочем', 'хорошо', 'свою', 'этой', 
            'перед', 'иногда', 'лучше', 'чуть', 'том', 'нельзя', 'такой', 'им', 
            'более', 'всегда', 'конечно', 'всю', 'между'
        }
        
    def extract_word_forms(self, text):
        """Извлечение словоформ (токенизация с сохранением исходных форм)"""
        # Приведение к нижнему регистру и разделение на слова
        words = re.findall(r'\b[а-яёa-z]+\b', text.lower())
        # Удаление стоп-слов
        words = [word for word in words if word not in self.stop_words and len(word) > 2]
        return words
    
    def extract_lemma_forms(self, text):
        """Извлечение начальных форм слов (простая лемматизация)"""
        words = self.extract_word_forms(text)
        lemmas = []
        
        for word in words:
            lemma = self.simple_lemmatize(word)
            lemmas.append(lemma)
            
        return lemmas
    
    def simple_lemmatize(self, word):
        """Простая лемматизация для русского языка"""
        if len(word) <= 3:
            return word
            
        # Правила для существительных (женский род)
        if word.endswith('ости') or word.endswith('асти'):
            return word[:-2] + 'ость'
        elif word.endswith('ации'):
            return word[:-3] + 'ация'
        elif word.endswith('ии'):
            return word[:-1]
            
        # Правила для прилагательных
        if word.endswith('ого') or word.endswith('его'):
            return word[:-3] + 'ий'
        elif word.endswith('ым') or word.endswith('им'):
            return word[:-2] + 'ый'
        elif word.endswith('ой') or word.endswith('ей'):
            return word[:-2] + 'ый'
            
        # Правила для глаголов
        if word.endswith('ить') or word.endswith('еть') or word.endswith('ать'):
            return word[:-2] + 'ь'
        elif word.endswith('ют') or word.endswith('ут'):
            return word[:-2]
        elif word.endswith('ят') or word.endswith('ат'):
            return word[:-2]
        elif word.endswith('ил') or word.endswith('ел'):
            return word[:-2] + 'ь'
            
        # Удаление окончаний множественного числа
        if word.endswith('ы') or word.endswith('и'):
            base = word[:-1]
            if len(base) > 2:
                return base
                
        return word
    
    def build_vocabulary(self, documents):
        """Построение словаря из всех документов"""
        all_words = []
        for doc in documents:
            all_words.extend(self.extract_word_forms(doc))
            all_words.extend(self.extract_lemma_forms(doc))
        
        self.vocabulary = set(all_words)
        self.word_to_index = {word: idx for idx, word in enumerate(sorted(self.vocabulary))}
        self.index_to_word = {idx: word for word, idx in self.word_to_index.items()}
    
    def create_word_vectors(self, documents, method='word_forms'):
        """Создание векторов документов"""
        vectors = []
        
        for doc in documents:
            if method == 'word_forms':
                words = self.extract_word_forms(doc)
            else:  # lemma_forms
                words = self.extract_lemma_forms(doc)
            
            # Создание вектора частот слов
            vector = [0] * len(self.vocabulary)
            word_counts = Counter(words)
            
            for word, count in word_counts.items():
                if word in self.word_to_index:
                    vector[self.word_to_index[word]] = count
            
            vectors.append(vector)
        
        return vectors

class DataProcessor:
    def __init__(self):
        self.vectorizer = DocumentVectorizer()
    
    def read_docx_file(self, file_path):
        """Чтение DOCX файла"""
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Игнорируем пустые параграфы
                    text.append(paragraph.text.strip())
            return '\n'.join(text)
        except Exception as e:
            print(f"Ошибка чтения DOCX файла {file_path}: {e}")
            return ""
    
    def read_xml_file(self, file_path):
        """Чтение XML файла с разметкой"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            return self.parse_xml_content(root)
        except Exception as e:
            print(f"Ошибка чтения XML файла {file_path}: {e}")
            return ""
    
    def parse_xml_content(self, root):
        """Парсинг XML контента"""
        text_parts = []
        
        for elem in root.iter():
            if elem.text and elem.text.strip():
                text_parts.append(elem.text.strip())
            if elem.tail and elem.tail.strip():
                text_parts.append(elem.tail.strip())
        
        return ' '.join(text_parts)
    
    def process_dataset(self, dataset_path):
        """Обработка всего набора данных"""
        documents = []
        document_names = []
        
        if os.path.isfile(dataset_path):
            # Одиночный файл
            if dataset_path.endswith('.docx'):
                text = self.read_docx_file(dataset_path)
                if text:
                    documents.append(text)
                    document_names.append(os.path.basename(dataset_path))
            elif dataset_path.endswith('.xml'):
                text = self.read_xml_file(dataset_path)
                if text:
                    documents.append(text)
                    document_names.append(os.path.basename(dataset_path))
        
        elif os.path.isdir(dataset_path):
            # Папка с файлами
            for file_name in os.listdir(dataset_path):
                file_path = os.path.join(dataset_path, file_name)
                if file_name.endswith('.docx'):
                    text = self.read_docx_file(file_path)
                    if text:
                        documents.append(text)
                        document_names.append(file_name)
                elif file_name.endswith('.xml'):
                    text = self.read_xml_file(file_path)
                    if text:
                        documents.append(text)
                        document_names.append(file_name)
        
        return documents, document_names
    
    def create_report(self, documents, document_names, output_path):
        """Создание отчёта в Excel"""
        # Построение словаря
        self.vectorizer.build_vocabulary(documents)
        
        # Создание векторов
        word_vectors = self.vectorizer.create_word_vectors(documents, 'word_forms')
        lemma_vectors = self.vectorizer.create_word_vectors(documents, 'lemma_forms')
        
        # Создание Excel файла
        wb = openpyxl.Workbook()
        
        # Лист с словоформами
        ws_word = wb.active
        ws_word.title = "Словоформы"
        
        # Заголовки
        headers = ['Документ'] + [self.vectorizer.index_to_word[i] for i in range(len(self.vectorizer.vocabulary))]
        ws_word.append(headers)
        
        # Данные
        for i, (doc_name, vector) in enumerate(zip(document_names, word_vectors)):
            row = [doc_name] + vector
            ws_word.append(row)
        
        # Лист с начальными формами
        ws_lemma = wb.create_sheet("Начальные_формы")
        ws_lemma.append(headers)
        
        for i, (doc_name, vector) in enumerate(zip(document_names, lemma_vectors)):
            row = [doc_name] + vector
            ws_lemma.append(row)
        
        # Лист со статистикой
        ws_stats = wb.create_sheet("Статистика")
        ws_stats.append(['Параметр', 'Значение'])
        ws_stats.append(['Всего документов', len(documents)])
        ws_stats.append(['Размер словаря', len(self.vectorizer.vocabulary)])
        
        # Подсчет общего количества слов
        total_word_forms = sum(len(self.vectorizer.extract_word_forms(doc)) for doc in documents)
        total_lemmas = sum(len(self.vectorizer.extract_lemma_forms(doc)) for doc in documents)
        
        ws_stats.append(['Общее количество слов (словоформы)', total_word_forms])
        ws_stats.append(['Общее количество слов (леммы)', total_lemmas])
        
        # Статистика по самому частому документу
        if documents:
            doc_lengths = [len(self.vectorizer.extract_word_forms(doc)) for doc in documents]
            ws_stats.append(['Максимум слов в документе', max(doc_lengths)])
            ws_stats.append(['Минимум слов в документе', min(doc_lengths)])
            ws_stats.append(['Среднее количество слов', sum(doc_lengths) // len(doc_lengths)])
        
        # Топ-10 самых частых слов
        all_words = []
        for doc in documents:
            all_words.extend(self.vectorizer.extract_word_forms(doc))
        
        word_freq = Counter(all_words)
        top_words = word_freq.most_common(10)
        
        ws_stats.append([])
        ws_stats.append(['Топ-10 самых частых слов:', 'Частота'])
        for word, freq in top_words:
            ws_stats.append([word, freq])
        
        # Сохранение
        wb.save(output_path)
        
        return word_vectors, lemma_vectors

def main():
    processor = DataProcessor()
    
    print("=== СИСТЕМА ВЕКТОРИЗАЦИИ ДОКУМЕНТОВ ===")
    
    # Чтение DOCX файла
    docx_file = "спортивная_медицина_01.docx"
    
    if os.path.exists(docx_file):
        print(f"Чтение файла: {docx_file}")
        document_text = processor.read_docx_file(docx_file)
        
        if document_text:
            print("Файл успешно прочитан!")
            print(f"Длина текста: {len(document_text)} символов")
            print(f"Количество абзацев: {len(document_text.split(chr(10)))}")
            
            # Показать первые 500 символов текста
            preview = document_text[:500] + "..." if len(document_text) > 500 else document_text
            print(f"\nПредварительный просмотр текста:\n{preview}")
            
            # Обработка документа
            documents = [document_text]
            document_names = [docx_file]
            
            # Демонстрация методов обработки
            print("\n=== АНАЛИЗ МЕТОДОВ ОБРАБОТКИ ===")
            
            # Словоформы
            word_forms = processor.vectorizer.extract_word_forms(document_text)
            print(f"Количество словоформ: {len(word_forms)}")
            print(f"Примеры словоформ: {word_forms[:15]}...")
            
            # Начальные формы
            lemma_forms = processor.vectorizer.extract_lemma_forms(document_text)
            print(f"Количество начальных форм: {len(lemma_forms)}")
            print(f"Примеры начальных форм: {lemma_forms[:15]}...")
            
            # Создание отчёта
            print("\n=== СОЗДАНИЕ ОТЧЁТА ===")
            output_file = "векторы_документов.xlsx"
            word_vectors, lemma_vectors = processor.create_report(documents, document_names, output_file)
            
            print(f"Отчёт сохранён в файл: {output_file}")
            print(f"Размер словаря: {len(processor.vectorizer.vocabulary)} слов")
            print(f"Размер вектора документа: {len(word_vectors[0])} измерений")
            
            # Показать некоторые статистики
            print(f"\n=== СТАТИСТИКА ===")
            total_unique_words = len(set(word_forms))
            total_unique_lemmas = len(set(lemma_forms))
            print(f"Уникальные словоформы: {total_unique_words}")
            print(f"Уникальные леммы: {total_unique_lemmas}")
            print(f"Сжатие словаря: {((total_unique_words - total_unique_lemmas) / total_unique_words * 100):.1f}%")
            
        else:
            print("Ошибка: не удалось прочитать файл или файл пуст")
    else:
        print(f"Файл {docx_file} не найден в текущей директории")
        print("Доступные файлы в текущей директории:")
        for file in os.listdir('.'):
            if file.endswith('.docx') or file.endswith('.xml'):
                print(f"  - {file}")

def create_technical_report():
    """Создание технического отчёта"""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Отчёт о реализации системы векторизации документов', 0)
    
    # Введение
    doc.add_heading('1. Введение', level=1)
    doc.add_paragraph(
        'Данный отчёт описывает реализацию системы для создания векторов документов '
        'с использованием двух методов обработки текста: выделения словоформ и '
        'выделения начальных форм слов. Система разработана для обработки размеченных '
        'экспертом наборов данных.'
    )
    
    # Алгоритмы обработки текста
    doc.add_heading('2. Реализованные алгоритмы', level=1)
    
    doc.add_heading('2.1. Выделение словоформ', level=2)
    doc.add_paragraph(
        'Алгоритм выделения словоформ включает следующие этапы:'
    )
    steps = [
        'Токенизация текста с использованием регулярных выражений',
        'Приведение всех слов к нижнему регистру',
        'Фильтрация стоп-слов (общеупотребительных слов)',
        'Удаление слов длиной менее 3 символов',
        'Сохранение исходных словоформ для построения векторов'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('2.2. Выделение начальных форм слов', level=2)
    doc.add_paragraph(
        'Алгоритм лемматизации реализует простые морфологические правила для русского языка:'
    )
    lemmatization_rules = [
        'Обработка существительных: удаление падежных окончаний',
        'Обработка прилагательных: приведение к именительному падежу',
        'Обработка глаголов: удаление личных окончаний',
        'Обработка окончаний множественного числа',
        'Сохранение коротких слов (менее 4 символов) без изменений'
    ]
    for rule in lemmatization_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    # Пример работы с файлом
    doc.add_heading('3. Пример обработки DOCX файла', level=1)
    
    # Попробуем прочитать и проанализировать файл
    processor = DataProcessor()
    if os.path.exists("спортивная_медицина_01.docx"):
        document_text = processor.read_docx_file("спортивная_медицина_01.docx")
        if document_text:
            word_forms = processor.vectorizer.extract_word_forms(document_text)
            lemma_forms = processor.vectorizer.extract_lemma_forms(document_text)
            
            doc.add_paragraph(f'Обработан файл: спортивная_медицина_01.docx')
            doc.add_paragraph(f'Размер текста: {len(document_text)} символов')
            doc.add_paragraph(f'Количество словоформ: {len(word_forms)}')
            doc.add_paragraph(f'Количество начальных форм: {len(lemma_forms)}')
            doc.add_paragraph(f'Уникальных словоформ: {len(set(word_forms))}')
            doc.add_paragraph(f'Уникальных лемм: {len(set(lemma_forms))}')
    
    doc.add_heading('4. Формат выходных данных', level=1)
    doc.add_paragraph(
        'Система генерирует Excel файл со следующими листами:'
    )
    sheets = [
        'Словоформы - векторы документов на основе словоформ',
        'Начальные_формы - векторы на основе лемм',
        'Статистика - общая статистика по набору данных'
    ]
    for sheet in sheets:
        doc.add_paragraph(sheet, style='List Bullet')
    
    doc.save('технический_отчёт.docx')
    print("Технический отчёт сохранён в файл 'технический_отчёт.docx'")


if __name__ == "__main__":

    main()
    create_technical_report()